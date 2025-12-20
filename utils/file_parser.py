# from docx import Document
# from pypdf import PdfReader

# def extract_text_from_file(uploaded_file) -> str:
#     file_type = uploaded_file.name.split(".")[-1].lower()

#     if file_type == "txt":
#         return uploaded_file.read().decode("utf-8")

#     elif file_type == "docx":
#         doc = Document(uploaded_file)
#         return "\n".join(p.text for p in doc.paragraphs)

#     elif file_type == "pdf":
#         reader = PdfReader(uploaded_file)
#         return "\n".join(page.extract_text() or "" for page in reader.pages)

#     else:
#         raise ValueError("Unsupported file format")

from docx import Document
from pypdf import PdfReader

def extract_text_from_file(uploaded_file) -> str:
    uploaded_file.seek(0)  # 🔥 CRITICAL

    file_type = uploaded_file.name.split(".")[-1].lower()

    if file_type == "txt":
        text = uploaded_file.read().decode("utf-8")

    elif file_type == "docx":
        doc = Document(uploaded_file)
        text = "\n".join(p.text for p in doc.paragraphs)

    elif file_type == "pdf":
        reader = PdfReader(uploaded_file)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)

    else:
        raise ValueError("Unsupported file format")

    uploaded_file.seek(0)  # 🔥 RESET AGAIN FOR SAFETY
    return text
